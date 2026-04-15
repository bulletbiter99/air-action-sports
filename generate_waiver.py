"""
Air Action Sport, LLC — Waiver Document Generator
Run this script in Claude Code to produce a properly formatted waiver.docx
Requires: pip install python-docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page setup ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1.25)
section.right_margin  = Inches(1.25)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ── Header ───────────────────────────────────────────────────────────────────
header = section.header
hp = header.paragraphs[0]
hp.text = "AIR ACTION SPORT, LLC — RELEASE OF LIABILITY AND ASSUMPTION OF RISK"
hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
hr = hp.runs[0]
hr.font.name = "Times New Roman"
hr.font.size = Pt(9)
hr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

# ── Footer with page numbers ─────────────────────────────────────────────────
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("Page ")
fr.font.name = "Times New Roman"
fr.font.size = Pt(9)
fld = OxmlElement("w:fldChar")
fld.set(qn("w:fldCharType"), "begin")
fp.runs[-1]._r.append(fld)
instr = OxmlElement("w:instrText")
instr.text = "PAGE"
fp.runs[-1]._r.append(instr)
fld2 = OxmlElement("w:fldChar")
fld2.set(qn("w:fldCharType"), "end")
fp.runs[-1]._r.append(fld2)
fr2 = fp.add_run(" | Document Version: 1.0 | Effective Date: April 2026")
fr2.font.name = "Times New Roman"
fr2.font.size = Pt(9)

# ── Styles ───────────────────────────────────────────────────────────────────
def set_font(run, bold=False, size=12, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic

def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    set_font(r, bold=True, size=13 if level == 1 else 12)
    return p

def add_subheading(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    set_font(r, bold=True, size=12)
    return p

def add_body(text, bold=False, caps=False, indent=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Inches(0.4)
    display = text.upper() if caps else text
    r = p.add_run(display)
    set_font(r, bold=bold)
    return p

def add_bullet(text, indent_level=1):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent  = Inches(0.4 * indent_level)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    set_font(r)
    return p

def add_sig_line(label, wide=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    width = 3.5 if wide else 2.5
    r = p.add_run(f"{label}:  " + "_" * int(width * 10))
    set_font(r)
    return p

def add_notice(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    r = p.add_run(text)
    set_font(r, bold=True, size=11)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ("top","left","bottom","right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"),  "6")
        b.set(qn("w:space"), "4")
        b.set(qn("w:color"), "444444")
        pBdr.append(b)
    pPr.append(pBdr)
    return p

def add_hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "888888")
    pBdr.append(bot)
    pPr.append(pBdr)

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        r = hdr_cells[i].paragraphs[0].runs[0]
        set_font(r, bold=True, size=11)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        cells = t.rows[ri+1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            r = cells[ci].paragraphs[0].runs[0]
            set_font(r, size=11)
    if col_widths:
        for i, row in enumerate(t.rows):
            for j, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[j])
    doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# DOCUMENT TITLE
# ════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("AIR ACTION SPORT, LLC")
set_font(r, bold=True, size=16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("RELEASE OF LIABILITY, ASSUMPTION OF RISK,\nAND PARTICIPANT AGREEMENT")
set_font(r, bold=True, size=13)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Document Version: 1.0  |  Effective Date: April 2026")
set_font(r, size=10, italic=True)
doc.add_paragraph()

add_hr()

# ── Instructions ─────────────────────────────────────────────────────────────
add_body(
    "INSTRUCTIONS: PLEASE READ THIS ENTIRE AGREEMENT CAREFULLY BEFORE SIGNING. "
    "THIS AGREEMENT CONTAINS A RELEASE OF LIABILITY, ASSUMPTION OF RISK, JURY TRIAL "
    "WAIVER, AND OTHER PROVISIONS THAT SIGNIFICANTLY AFFECT YOUR LEGAL RIGHTS. BY "
    "SIGNING THIS AGREEMENT, YOU ARE, AMONG OTHER THINGS, EXPRESSLY WAIVING YOUR "
    "RIGHT TO SUE OR SEEK MONEY DAMAGES FROM AIR ACTION SPORT, LLC ('AAS') AND ITS "
    "MANAGERS, MEMBERS, EMPLOYEES, INDEPENDENT CONTRACTORS, AGENTS, AND AFFILIATES "
    "FOR NEGLIGENCE IF A PERSONAL INJURY IS SUSTAINED AT ANY AAS PREMISES OR EVENT. "
    "If you do not agree to any term, provision, or paragraph of this Agreement, DO "
    "NOT SIGN the Agreement and please exit the Premises immediately.",
    bold=True
)
add_hr()

# ── Definitions ──────────────────────────────────────────────────────────────
add_heading("DEFINITIONS", 2)

defs = [
    ('"AAS"',
     "means Air Action Sport, LLC, a Utah limited liability company, and all of its managers, "
     "members, employees, agents, officers, directors, affiliates, volunteers, participants, "
     "clients, customers, invitees, independent contractors, insurers, facility operators, and "
     "landowners, together with their respective successors and assigns (collectively, "
     '"The Released And Indemnified Parties").'),
    ('"Premises"',
     "means any and all locations, properties, sites, fields, structures, parking areas, camping "
     "zones, staging areas, vendor areas, and surrounding land used by AAS in connection with its "
     "events and activities, whether now existing or established in the future, including but not "
     "limited to all current operating sites listed in the Site Schedule attached hereto as Exhibit A, "
     "any temporary or pop-up event locations, and any privately leased or licensed properties used "
     "by AAS for hosted events. The term 'Premises' shall apply regardless of whether a specific "
     "site address is listed in Exhibit A at the time this Agreement is signed."),
    ('"Activities"',
     "means airsoft gaming events, milsim operations, skirmish sessions, scenario games, night "
     "operations, private hire events, camping, spectating, and any other recreational or related "
     "activities conducted at the Premises, including travel to and from any area of the Premises."),
    ('"Claim Period"',
     "means the period beginning on the Effective Date and ending at midnight 365 days after the "
     "Effective Date."),
    ('"Effective Date"',
     "means the date and time this Agreement is electronically or physically signed by the Participant."),
    ('"Releasing Parties"',
     "means the signing participant, their spouse, heir(s), personal representative(s), and their "
     "respective successors and assigns."),
]

for term, definition in defs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(5)
    r1 = p.add_run(term + "  ")
    set_font(r1, bold=True)
    r2 = p.add_run(definition)
    set_font(r2)

add_hr()

# ── Preamble ─────────────────────────────────────────────────────────────────
add_heading("PARTICIPANT AGREEMENT AND RELEASE OF LIABILITY", 2)
add_body(
    "All participants in Activities on the Premises must complete and sign this Agreement prior "
    "to participation. This Agreement is effective as of the Effective Date and shall remain valid "
    "and effective to release and indemnify AAS from any Claims arising during the Claim Period."
)

# ── Sections 1-22 ────────────────────────────────────────────────────────────
clauses = [
    ("1. RELEASE AND INDEMNITY",
     [("body",
       "For myself and on behalf of the Releasing Parties, I hereby agree to release, remise, "
       "forever discharge, defend, hold harmless, and indemnify The Released And Indemnified "
       "Parties from and against any and all claims, actions, causes of action, proceedings, "
       "suits, costs, liabilities, damages, and expenses, whether known or unknown (including "
       "but not limited to all direct, special, incidental, exemplary, punitive, and "
       "consequential damages, losses of any kind and attorney fees), and however caused, "
       "including without limitation by negligent conduct (hereafter collectively, 'Claims') "
       "of any and all of the Releasing Parties that arise on, are based upon, or result from, "
       "any act, event, occurrence, or omission on the Premises during the Claim Period."),
      ("body",
       "I agree not to initiate, prosecute, or aid any other party in prosecuting any Claim of "
       "any kind against The Released And Indemnified Parties arising from the Activities, "
       "whether based on common law, equity, or any federal, state, or local statute, ordinance, "
       "or rule of law. This release expressly includes claims arising from the ordinary "
       "negligence of AAS and The Released And Indemnified Parties, but does not extend to "
       "claims arising from gross negligence or willful misconduct by AAS.")]),

    ("2. ACKNOWLEDGEMENT OF RISKS",
     [("body",
       "I acknowledge that airsoft and related Activities at the Premises are inherently "
       "dangerous and hazardous by their very nature. By participating in, observing, or "
       "allowing participation in the Activities, I expressly assume all risks associated with "
       "the Activities on behalf of myself and the Releasing Parties and expressly contract not "
       "to sue for any injury or illness sustained as a result of such participation. Known and "
       "foreseeable risks include but are not limited to:"),
      ("bullets", [
          "Cuts, bruises, welts, burns, and abrasions from projectile impact",
          "Tripping, falling, spraining, or breaking bones including wrists, ankles, necks, backs, and skulls",
          "Eye and facial injuries including from projectiles, debris, and environmental hazards",
          "Head trauma, concussion, muscle pulls, exhaustion, and permanent paralysis",
          "Exposure to bright, colored, or strobe lighting which may induce seizure or disorientation",
          "Collisions with other participants, spectators, fixed objects, structures, and terrain features",
          "Night operations and low-light scenario hazards including reduced visibility, disorientation, and trip hazards",
          "Weather-related risks including extreme heat, cold, lightning, high winds, flooding, and sudden weather changes",
          "Environmental hazards including uneven terrain, wildlife, insects, thorns, dust, allergens, and bodies of water",
          "Infectious disease transmission due to the public nature of the Activities",
          "Risks arising from the actions of third parties, other participants, and spectators",
          "Other serious injuries, permanent disability, or death",
      ]),
      ("body",
       "I understand that no matter how carefully AAS manages the Activities, the risk of "
       "serious injury and illness is not eliminated and remains foreseeable.")]),

    ("3. ASSUMPTION OF RISK AND LOSS",
     [("body_caps",
       "I knowingly and freely assume all known and unknown risks of injury, illness, damage, "
       "and/or death on behalf of myself and the Releasing Parties. My participation is purely "
       "voluntary. I agree to pay for the cost of any medical assistance requested by AAS on "
       "behalf of any Releasing Party and assume full financial responsibility for any damage, "
       "illness, or injury occurring on the Premises. I further assume the risk of aggravation "
       "of any preexisting medical or physical condition, whether known or unknown.")]),

    ("4. PERSONAL PROTECTIVE EQUIPMENT (PPE) COMPLIANCE",
     [("body", "I acknowledge and agree that:"),
      ("bullets", [
          "(a) Full-seal eye protection meeting ANSI Z87.1 or equivalent standard is mandatory at all times on active playing fields. I agree to wear all required PPE throughout gameplay.",
          "(b) AAS and its marshals have the right and authority to inspect my PPE and deny my access to the playing field if PPE is deemed non-compliant or inadequate.",
          "(c) Removing eye protection on an active field is grounds for immediate removal from the event without refund.",
          "(d) I am solely responsible for ensuring that any minor in my care wears appropriate PPE at all times.",
      ])]),

    ("5. VELOCITY (FPS) AND EQUIPMENT INSPECTION",
     [("body", "I acknowledge and agree that:"),
      ("bullets", [
          "(a) AAS enforces velocity limits on all airsoft equipment. I consent to having my equipment chronographed and inspected by AAS staff or marshals before and during any event.",
          "(b) Any firearm or replica exceeding the posted FPS limit will be prohibited from use for the duration of the event. AAS will secure non-compliant equipment and return it to the participant upon departure. AAS assumes no liability for the condition of equipment that was non-compliant at the time of inspection.",
          "(c) I represent that all equipment I bring to the Premises is legal to possess under applicable federal, state, and local law.",
          "(d) I agree not to modify equipment on-site to circumvent velocity limits.",
      ])]),

    ("6. WEATHER AND ENVIRONMENTAL CONDITIONS",
     [("body", "I acknowledge that outdoor airsoft events are subject to weather and environmental conditions beyond AAS's control. I understand and agree that:"),
      ("bullets", [
          "(a) AAS may modify, delay, or cancel events due to weather at its sole discretion.",
          "(b) I am responsible for my own physical preparation for outdoor conditions including heat, cold, sun exposure, and terrain.",
          "(c) I will follow all AAS protocols for weather-related emergencies including lightning protocols, evacuation procedures, and shelter-in-place instructions.",
          "(d) I release AAS from liability for injuries or discomfort arising from weather and environmental conditions.",
      ])]),

    ("7. OVERNIGHT CAMPING",
     [("body", "If I have purchased a camping add-on or am otherwise permitted to camp overnight at the Premises, I additionally acknowledge and agree that:"),
      ("bullets", [
          "(a) Overnight camping involves additional risks including fire hazards, wildlife encounters, carbon monoxide risks, cold exposure, and trip hazards in darkness.",
          "(b) I will comply with all posted camping rules including quiet hours, fire restrictions, and prohibited zones.",
          "(c) I will not enter any active playing field or restricted area during overnight hours.",
          "(d) AAS is not responsible for the security of personal property in the camping zone.",
          "(e) I am responsible for properly extinguishing any fire before sleeping or leaving the camping area.",
      ])]),

    ("8. INJURIES BY AND TO THIRD PARTIES",
     [("body",
       "I acknowledge that the Releasing Parties may be injured by the actions of other "
       "customers or invitees of AAS at the Premises ('Third Parties'). I agree to release, "
       "discharge, waive, defend, and indemnify The Released And Indemnified Parties against "
       "any Claims arising from acts or omissions of Third Parties during the Claim Period. "
       "I also acknowledge that acts or omissions of the Releasing Parties may cause injury "
       "to others, and agree to defend and indemnify The Released And Indemnified Parties and "
       "any third party against any Claim caused in whole or in part by the Releasing Parties.")]),

    ("9. SPONSOR AND VENDOR ZONE ACKNOWLEDGMENT",
     [("body", "I acknowledge that AAS events may include sponsor booths, vendor tents, product demonstrations, and third-party commercial activity on or adjacent to the Premises. I agree that:"),
      ("bullets", [
          "(a) AAS is not responsible for the acts, omissions, products, or representations of any sponsor or vendor.",
          "(b) I release AAS from any Claims arising from my interaction with sponsor or vendor activities.",
          "(c) I will comply with all safety requirements in sponsor and vendor zones as directed by AAS staff.",
      ])]),

    ("10. INSURANCE",
     [("body",
       "I certify that I have adequate personal insurance or sufficient personal assets to "
       "fully indemnify The Released And Indemnified Parties against any Claims for which I "
       "have an indemnity obligation under this Agreement, including Claims by third parties "
       "caused in whole or in part by my acts or omissions.")]),

    ("11. RULES AND SAFETY STANDARDS",
     [("body",
       "I acknowledge that I have read, understand, and agree to abide by all posted and "
       "presented rules and safety standards for the Activities at the Premises, including "
       "field-specific rules, marshal instructions, and any rules communicated at pre-game "
       "safety briefings. I acknowledge that failure to comply with rules may result in "
       "immediate removal from the event without refund. AAS's full Cancellation and Refund "
       "Policy is posted at airactionsport.com and is incorporated by reference into this "
       "Agreement.")]),

    ("12. REPRESENTATIONS AND PHYSICAL CONDITION",
     [("body",
       "I represent that I am physically able to participate in the Activities and have no "
       "preexisting physical or medical condition, including allergies, exercise-induced "
       "conditions, heart conditions, seizure disorders, or conditions induced by strobe or "
       "low lighting, that would endanger me during the Activities. I represent that I will "
       "conduct myself in a safe and responsible manner so as not to endanger the lives or "
       "property of any persons on the Premises.")]),

    ("13. BASIS OF BARGAIN",
     [("body",
       "I understand that AAS would not allow use of the Premises without my agreement to "
       "the terms and conditions herein.")]),

    ("14. CHOICE OF LAW AND VENUE",
     [("body",
       "This Agreement shall be governed by and construed under the laws of the State of "
       "Utah, without regard to conflicts of law principles. Venue for any dispute shall be "
       "exclusively in the Second Judicial District Court, Davis County, Utah. I agree to "
       "indemnify and hold The Released And Indemnified Parties harmless for all attorney "
       "fees and costs incurred in enforcing this Agreement.")]),

    ("15. PHOTOGRAPHY, VIDEO, AND DRONE POLICY",
     [("body", "I acknowledge and agree that:"),
      ("bullets", [
          "(a) AAS and its authorized media partners may photograph or video record events on the Premises. I grant AAS the irrevocable right to use my name, face, likeness, voice, and appearance in connection with exhibitions, publicity, advertising, and promotional materials without compensation or limitation.",
          "(b) I will not operate any drone or unmanned aerial vehicle on or over the Premises without prior written approval from AAS management obtained no less than 48 hours prior to the event.",
          "(c) I will not photograph or video record other participants without their consent in a manner that violates their reasonable expectation of privacy.",
      ])]),

    ("16. SOCIAL MEDIA RELEASE",
     [("body",
       "I acknowledge that I may post content related to AAS events on social media "
       "platforms. I agree not to post content that defames AAS, its staff, or other "
       "participants, or that depicts safety violations or rule infractions in a manner "
       "designed to harm AAS's reputation. AAS reserves the right to request removal of "
       "content that violates posted rules or applicable law.")]),

    ("17. MEDICAL EMERGENCY AUTHORIZATION",
     [("body", "In the event I am incapacitated and unable to communicate, I authorize AAS staff and marshals to:"),
      ("bullets", [
          "(a) Call 911 and emergency medical services on my behalf",
          "(b) Administer basic first aid",
          "(c) Provide responding emergency personnel with any medical information I have disclosed in my participant profile or booking record",
      ]),
      ("body",
       "I understand that AAS staff are not medical professionals and release AAS from "
       "liability for any first aid rendered in good faith.")]),

    ("18. DATA PRIVACY",
     [("body",
       "I acknowledge and consent to AAS collecting and storing my personal information "
       "(name, date of birth, contact details, emergency contact, and signed waiver) for the "
       "purpose of managing event participation, safety records, and legal compliance. AAS "
       "will not sell my personal information to third parties. My data will be retained for "
       "a minimum of 7 years following the Claim Period for adult participants. Records "
       "involving participants who were under 18 at the time of participation will be "
       "retained until that participant's 23rd birthday.")]),

    ("19. INDEMNITY",
     [("body",
       "In addition to and not in lieu of other indemnity provisions herein, I agree on "
       "behalf of myself and my spouse to indemnify and hold harmless AAS and all Released "
       "And Indemnified Parties from and against any and all losses, liabilities, claims, "
       "obligations, costs, damages, and expenses, including attorney fees, directly or "
       "indirectly arising out of my or my spouse's acts or omissions while participating in "
       "Activities at the Premises, unless it is determined that such liability resulted from "
       "the gross negligence or willful misconduct of AAS.")]),

    ("20. MISCELLANEOUS / SEVERABILITY",
     [("body",
       "This Agreement is intended to be as broad and inclusive as permitted by Utah law. "
       "If any clause is determined to be unenforceable, it shall be severed and the "
       "remainder shall continue in full legal force and effect. This Agreement represents "
       "the entire understanding of the parties. No subsequent modification is binding unless "
       "reduced to writing and signed by both parties.")]),

    ("21. ANNUAL RENEWAL",
     [("body",
       "This Agreement is valid for the Claim Period (365 days from the Effective Date). "
       "Participants who return to AAS events after expiration of their Claim Period are "
       "required to execute a new Agreement prior to participation.")]),

    ("22. JURY TRIAL WAIVER",
     [("body",
       "I, on behalf of myself and the Releasing Parties, hereby waive to the full extent "
       "permitted by applicable law any right to trial by jury in any proceeding arising out "
       "of or relating to this Agreement, the Activities, or any injury sustained in "
       "connection with the Activities. I represent that no party has represented that this "
       "waiver would not be enforced, and that all parties have been induced to enter this "
       "Agreement in part by this jury trial waiver.")]),
]

for title, items in clauses:
    add_subheading(title)
    for item in items:
        if item[0] == "body":
            add_body(item[1])
        elif item[0] == "body_caps":
            add_body(item[1], caps=True, bold=True)
        elif item[0] == "bullets":
            for b in item[1]:
                add_bullet(b)

# Jury waiver initial line
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(10)
r = p.add_run("Participant Initials acknowledging Jury Trial Waiver:  __________     Date:  __________")
set_font(r)

add_hr()

# Closing acknowledgment
add_body(
    "I HAVE READ THIS RELEASE AGREEMENT IN ITS ENTIRETY, FULLY UNDERSTAND ITS TERMS, "
    "UNDERSTAND THAT I AM GIVING UP SUBSTANTIAL LEGAL RIGHTS BY SIGNING IT, INCLUDING "
    "THE RIGHT TO SUE FOR NEGLIGENCE, AND SIGN IT FREELY AND VOLUNTARILY WITHOUT ANY "
    "INDUCEMENT.",
    bold=True, caps=True
)

add_hr()
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# ADULT SIGNATURE BLOCK
# ════════════════════════════════════════════════════════════════════════════
add_heading("ADULT PARTICIPANT SIGNATURE BLOCK", 2)
add_body("By signing below I confirm I am 18 years of age or older.")
doc.add_paragraph()

sig_fields = [
    "Participant Full Legal Name",
    "Signature",
    "Date",
    "Date of Birth",
    "Email Address",
    "Phone Number",
    "Emergency Contact Name",
    "Emergency Contact Phone",
]
for f in sig_fields:
    add_sig_line(f)

add_body("Known Medical Conditions / Allergies (optional but encouraged):  " + "_" * 50)
add_hr()
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# MINOR PARTICIPANT SECTION
# ════════════════════════════════════════════════════════════════════════════
add_heading("MINOR PARTICIPANT SECTION", 2)

add_notice(
    "NOTICE: Under Hawkins v. Peart, 37 P.3d 1062 (Utah 2001), reaffirmed in Rutherford v. "
    "Talisker Canyons Finance Co., 445 P.3d 474 (Utah 2019), parental pre-injury liability "
    "waivers and indemnity provisions on behalf of minor children are unenforceable as against "
    "public policy in Utah. This section does NOT constitute a valid liability release for "
    "minor participants and must not be relied upon as such. AAS must maintain adequate "
    "liability insurance explicitly covering minor participants before allowing any minor to "
    "participate. This section serves as parental consent, emergency medical authorization, "
    "media release, and age-tier acknowledgment only."
)

add_subheading("Age Participation Policy")
add_table(
    headers=["Age", "Requirement"],
    rows=[
        ["Under 12",
         "Not permitted to participate in any AAS event under any circumstances."],
        ["12\u201315",
         "Permitted with parent or legal guardian physically present and supervising "
         "on-site for the full duration of the event. Remote or off-site parental "
         "consent is not sufficient for this age group."],
        ["16\u201317",
         "Permitted with prior written parental or legal guardian approval. This signed "
         "form constitutes that approval. Supervising adult does not need to be present "
         "on-site."],
        ["18+",
         "No parental consent required. Adult participant signs independently."],
    ],
    col_widths=[1.0, 4.5]
)

add_body(
    "AAS reserves the right to require government-issued photo ID to verify participant age "
    "at check-in. Participants who cannot provide age verification may be denied entry. "
    "AAS reserves the right to deny entry to any participant who does not meet the applicable "
    "age requirement or whose supervising adult is not present as required."
)

add_subheading("Parental / Guardian Consent and Acknowledgment")
add_body(
    "I represent that I am the parent, legal guardian, or authorized custodian of the Minor "
    "participant listed below. I have read and understand this entire Agreement in full and "
    "acknowledge that all risk disclosures and Activity descriptions herein apply to my "
    "Minor's participation. I consent to my Minor's participation in the Activities at any "
    "AAS Premises. I additionally grant AAS the irrevocable right to photograph, record, "
    "and use the Minor's name, face, likeness, voice, and appearance in connection with "
    "exhibitions, publicity, advertising, and promotional materials without compensation or "
    "limitation. I acknowledge that I have reviewed the Age Participation Policy above and "
    "confirm that my Minor meets the applicable age requirement for the event(s) they are "
    "attending."
)
doc.add_paragraph()

minor_sigs = [
    "Minor's Full Legal Name",
    "Minor's Date of Birth",
    "Minor's Age at Time of Signing",
    "Parent / Guardian Full Legal Name",
    "Parent / Guardian Signature",
    "Date",
    "Relationship to Minor",
    "Parent / Guardian Phone (Day of Event)",
    "Emergency Contact Name (if different)",
    "Emergency Contact Phone",
]
for f in minor_sigs:
    add_sig_line(f)
add_body("Minor's Known Medical Conditions / Allergies:  " + "_" * 50)
doc.add_paragraph()
add_body("Parent / Guardian Initials acknowledging Notice and Age Policy:  __________")

add_subheading("On-Site Supervising Adult Acknowledgment")
add_body("(Required only if Minor is age 12\u201315. If Minor is age 16\u201317, skip this section.)")
add_body(
    "I, the undersigned, confirm that I will be physically present on-site at the AAS "
    "Premises for the full duration of the event in which the Minor named above is "
    "participating. I understand that:"
)
add_bullet("(a) I am personally responsible for the supervision and conduct of the Minor for the entirety of the event.")
add_bullet("(b) I will not leave the Premises while the Minor is participating without first notifying AAS staff and arranging for another qualified adult to assume supervision.")
add_bullet("(c) I will ensure the Minor wears all required PPE at all times on active fields and complies with all AAS rules and marshal instructions.")
add_bullet("(d) I have read and understood this entire Agreement and acknowledge the risks described herein.")
doc.add_paragraph()

super_sigs = [
    "Supervising Adult Full Legal Name",
    "Supervising Adult Signature",
    "Date",
    "Relationship to Minor",
    "Supervising Adult Phone (Day of Event)",
]
for f in super_sigs:
    add_sig_line(f)

add_body("Is the Supervising Adult the same person as the Parent / Guardian above?   Yes  /  No  (circle one)")
add_body("If No \u2014 the Parent / Guardian must also complete the Parental / Guardian Consent section above.")

add_hr()
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# ELECTRONIC SIGNATURE ACKNOWLEDGMENT
# ════════════════════════════════════════════════════════════════════════════
add_heading("ELECTRONIC SIGNATURE ACKNOWLEDGMENT", 2)
add_body(
    "I HEREBY ACKNOWLEDGE (1) THAT THIS DOCUMENT IS ELECTRONICALLY SIGNED IN ACCORDANCE "
    "WITH UTAH CODE ANN. \u00a7 46-4-201 AND (2) THAT THIS DOCUMENT IS VALID AND MAY BE "
    "ENFORCED IN THE SAME MANNER AS A HAND-SIGNED DOCUMENT. I ACKNOWLEDGE THE VALIDITY "
    "OF MY ELECTRONIC SIGNATURE AND WAIVE ANY RIGHT TO CLAIM THIS DOCUMENT IS INVALID OR "
    "UNENFORCEABLE BASED ON ITS ELECTRONIC FORM OR ELECTRONIC SIGNATURE.",
    bold=True, caps=True
)

add_hr()
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# EXHIBIT A
# ════════════════════════════════════════════════════════════════════════════
add_heading("EXHIBIT A \u2014 SITE SCHEDULE", 1)

add_body(
    "The following sites are current operating locations of Air Action Sport, LLC. GPS "
    "coordinates are provided in place of physical addresses where no formal street address "
    "exists, consistent with standard practice for outdoor recreational venues. This list "
    "may be updated from time to time. The definition of 'Premises' in this Agreement "
    "covers all sites listed below as well as any future sites, temporary locations, or "
    "privately leased properties used by AAS, regardless of whether listed here. Sites "
    "marked Coming Soon are included in the scope of this Agreement as of the date they "
    "become operational."
)

add_table(
    headers=["Site #", "Site Name", "Description", "GPS Coordinates", "Status"],
    rows=[
        ["01", "Delta Base",      "Woodland Site",              "[INSERT COORDINATES]", "Active"],
        ["02", "Trench Warfare",  "CQB \u2014 Echo Urban Warehouse", "[INSERT COORDINATES]", "Active"],
        ["03", "Foxtrot Fields",  "Open Field Site",            "[INSERT COORDINATES]", "Coming Soon"],
        ["[04]", "[Future Site]", "[Description]",              "[INSERT COORDINATES]", "[Status]"],
    ],
    col_widths=[0.6, 1.4, 1.8, 1.8, 1.0]
)

add_body(
    "Additional sites will be added to this Exhibit as AAS expands operations. The Agreement "
    "remains valid and applicable at all sites regardless of the date a site is added to "
    "this Schedule."
)

# ── Save ─────────────────────────────────────────────────────────────────────
output_path = "C:/Users/bulle/OneDrive/Desktop/Claude Code Projects/action-air-sports/AAS_Release_of_Liability_v1.0.docx"
doc.save(output_path)
print(f"Done! Saved: {output_path}")
